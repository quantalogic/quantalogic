"""Tool for converting markdown content to well-structured DOCX documents.

Why this tool:
- Provides a standardized way to convert markdown to professional DOCX documents
- Maintains consistent styling and formatting across documents
- Handles complex elements like diagrams, code blocks, and tables
- Supports customization through templates and style configurations
"""

import hashlib
import json
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import markdown
import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from loguru import logger
from PIL import Image
from pygments import highlight
from pygments.formatters import HtmlFormatter
from pygments.lexers import TextLexer, get_lexer_by_name
from pygments.styles import get_style_by_name

from quantalogic.tools.tool import Tool, ToolArgument


class MarkdownToDocxTool(Tool):
    """Converts markdown to professional DOCX documents with advanced formatting."""

    model_config = {
        "arbitrary_types_allowed": True
    }
    
    name: str = "markdown_to_docx_tool"
    description: str = (
        "Converts markdown to DOCX with support for images, Mermaid diagrams, "
        "code blocks, and advanced document formatting."
    )
    need_validation: bool = False
    
    arguments: List[ToolArgument] = [
        ToolArgument(
            name="markdown_content",
            arg_type="string",
            description="Markdown content with support for advanced formatting",
            required=True,
            example='''# Technical Documentation

## Overview
This document demonstrates various formatting capabilities.

## Code Examples

### Python Code
```python
def hello_world():
    """Greet the world."""
    return "Hello, World!"
```

### JavaScript Code
```javascript
function calculateTotal(items) {
    return items.reduce((sum, item) => sum + item.price, 0);
}
```

## System Architecture
```mermaid
graph TD
    A[Frontend] --> B[API Gateway]
    B --> C[Microservices]
    C --> D[(Database)]
    B --> E[Cache]
```

## Feature List
1. **Authentication**
   - OAuth 2.0 support
   - Multi-factor authentication
   - Role-based access control

2. **Data Processing**
   - Real-time analytics
   - Batch processing
   - Data validation

## Performance Metrics
| Metric | Value | Status |
|--------|--------|--------|
| Latency | 100ms | ✅ |
| Uptime | 99.9% | ✅ |
| Error Rate | 0.1% | ✅ |

> **Note**: All metrics are measured over a 30-day period.

![System Dashboard](https://example.com/dashboard.png)
''',
        ),
        ToolArgument(
            name="output_path",
            arg_type="string",
            description="Path for saving the DOCX file",
            required=True,
            example="/path/to/output.docx",
        ),
        ToolArgument(
            name="template_path",
            arg_type="string",
            description="Optional DOCX template path. Use a template for consistent corporate styling.",
            required=False,
            example="/path/to/template.docx",
        ),
        ToolArgument(
            name="style_config",
            arg_type="string",
            description="JSON string with style settings",
            required=False,
            example='''{
    "font_name": "Arial",
    "title_size": 32,
    "heading1_size": 28,
    "heading2_size": 24,
    "heading3_size": 20,
    "body_size": 11,
    "code_size": 10,
    "code_font": "Consolas",
    "primary_color": [0, 112, 192],
    "secondary_color": [68, 114, 196],
    "text_color": [0, 0, 0],
    "link_color": [0, 0, 255],
    "code_background": [245, 245, 245],
    "code_border_color": [200, 200, 200],
    "table_header_background": [217, 217, 217],
    "margins": {
        "top": 1,
        "bottom": 1,
        "left": 1,
        "right": 1
    }
}''',
        ),
    ]

    # Default style configuration
    DEFAULT_STYLES: Dict[str, Union[str, int, List[int]]] = {
        "font_name": "Calibri",
        "title_size": 32,
        "heading1_size": 28,
        "heading2_size": 24,
        "heading3_size": 20,
        "body_size": 11,
        "code_size": 10,
        "code_font": "Consolas",
        "primary_color": [0, 112, 192],
        "secondary_color": [68, 114, 196],
        "text_color": [0, 0, 0],
        "link_color": [0, 0, 255],
        "code_background": [245, 245, 245],
        "code_border_color": [200, 200, 200],
        "table_header_background": [217, 217, 217],
        "margins": {
            "top": 1,
            "bottom": 1,
            "left": 1,
            "right": 1
        }
    }

    def _normalize_path(self, path: str) -> Path:
        """Convert path string to normalized Path object.
        
        Args:
            path: Input path string
            
        Returns:
            Normalized Path object
        """
        if path.startswith("~"):
            path = os.path.expanduser(path)
        return Path(path).resolve()

    def _parse_style_config(self, style_config: Optional[str]) -> Dict:
        """Parse and validate style configuration.
        
        Args:
            style_config: JSON style configuration string
            
        Returns:
            Merged style configuration dictionary
        """
        config = self.DEFAULT_STYLES.copy()
        if style_config:
            try:
                custom_styles = json.loads(style_config)
                config.update(custom_styles)
            except json.JSONDecodeError as e:
                logger.warning(f"Invalid style config, using defaults: {e}")
        return config

    def _create_document(self, template_path: Optional[str] = None) -> Document:
        """Create a new document with predefined styles."""
        if template_path:
            template_path = self._normalize_path(template_path)
            if not template_path.exists():
                logger.warning(f"Template not found: {template_path}. Using default template.")
                doc = Document()
            else:
                doc = Document(template_path)
        else:
            doc = Document()

        # Add custom styles if they don't exist
        styles = doc.styles

        # Code block style
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = self.DEFAULT_STYLES["code_font"]
            code_font.size = Pt(self.DEFAULT_STYLES["code_size"])
            code_style.paragraph_format.space_before = Pt(12)
            code_style.paragraph_format.space_after = Pt(12)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.right_indent = Inches(0.5)
            code_style.paragraph_format.first_line_indent = Inches(0)
            
        # Code header style
        if 'CodeHeader' not in styles:
            header_style = styles.add_style('CodeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = self.DEFAULT_STYLES["font_name"]
            header_font.size = Pt(self.DEFAULT_STYLES["code_size"] + 2)
            header_font.bold = True
            header_style.paragraph_format.space_before = Pt(12)
            header_style.paragraph_format.space_after = Pt(6)
            header_style.paragraph_format.left_indent = Inches(0.5)
            
        # Diagram style
        if 'Diagram' not in styles:
            diagram_style = styles.add_style('Diagram', WD_STYLE_TYPE.PARAGRAPH)
            diagram_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            diagram_style.paragraph_format.space_before = Pt(12)
            diagram_style.paragraph_format.space_after = Pt(12)
            
        # Caption style
        if 'Caption' not in styles:
            caption_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
            caption_font = caption_style.font
            caption_font.name = self.DEFAULT_STYLES["font_name"]
            caption_font.size = Pt(10)
            caption_font.italic = True
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_after = Pt(18)

        return doc

    def _apply_text_style(self, run, style: str, style_config: Dict) -> None:
        """Apply text styling to a run.
        
        Args:
            run: Document run to style
            style: Style to apply ('bold', 'italic', 'code', or 'link')
            style_config: Style configuration dictionary
        """
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.name = style_config["code_font"]
        elif style == 'link':
            run.font.color.rgb = RGBColor(*style_config["link_color"])
            run.underline = True

    def _handle_image(self, src: str) -> Optional[str]:
        """Process and save image from source.
        
        Args:
            src: Image source (URL or path)
            
        Returns:
            Path to processed image or None if failed
        """
        try:
            if src.startswith(('http://', 'https://')):
                response = requests.get(src)
                response.raise_for_status()
                path = f"image_{hash(src)}.{src.split('.')[-1]}"
                with open(path, 'wb') as f:
                    f.write(response.content)
                return path
            return src
        except Exception as e:
            logger.error(f"Failed to process image {src}: {e}")
            return None

    def _add_image_to_doc(self, doc: Document, image_path: str) -> None:
        """Add image to document with proper sizing.
        
        Args:
            doc: Target document
            image_path: Path to image file
        """
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            img = Image.open(image_path)
            width = min(6.0, img.width / 96)  # Max 6 inches, convert from pixels
            doc.add_picture(image_path, width=Inches(width))
        except Exception as e:
            logger.error(f"Failed to add image {image_path}: {e}")

    def _process_mermaid(self, code: str) -> Tuple[Optional[str], Optional[str]]:
        """Generate image from Mermaid diagram code.
        
        Args:
            code: Mermaid diagram source code
            
        Returns:
            Tuple of (diagram_path, error_message)
        """
        try:
            # Create hash of code for caching
            code_hash = hashlib.md5(code.encode()).hexdigest()
            cache_dir = Path(tempfile.gettempdir()) / "mermaid_cache"
            cache_dir.mkdir(exist_ok=True)
            
            cache_path = cache_dir / f"{code_hash}.png"
            if cache_path.exists():
                return str(cache_path), None
            
            # Use Mermaid.ink API for rendering
            import base64
            graphbytes = base64.b64encode(code.encode('utf-8'))
            graphurl = f"https://mermaid.ink/img/{graphbytes.decode('utf-8')}"
            
            response = requests.get(graphurl)
            if response.status_code == 200:
                with open(cache_path, "wb") as f:
                    f.write(response.content)
                return str(cache_path), None
            else:
                error = f"Failed to generate Mermaid diagram: HTTP {response.status_code}"
                logger.error(error)
                return None, error
                
        except Exception as e:
            error = f"Failed to generate diagram: {e}"
            logger.error(error)
            return None, error

    def _process_code_block(self, doc: Document, element: Tag, style_config: Dict):
        """Process a code block element and add it to document with syntax highlighting."""
        try:
            # Get language if specified
            code_class = element.get('class', [])
            language = code_class[0].replace('language-', '') if code_class else 'text'
            
            # Get code content preserving whitespace
            code_text = element.get_text().strip()
            if not code_text:
                return
                
            # Add language header
            header = doc.add_paragraph(style='CodeHeader')
            lang_run = header.add_run(f"{language.upper()}")
            lang_run.font.color.rgb = RGBColor(*style_config["secondary_color"])
            
            # Create code block container
            code_para = doc.add_paragraph(style='Code')
            
            # Apply syntax highlighting
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except:
                lexer = TextLexer()
                
            formatter = HtmlFormatter(
                style=get_style_by_name('monokai'),
                linenos=True,
                cssclass="source",
                linenostart=1
            )
            
            highlighted = highlight(code_text, lexer, formatter)
            soup = BeautifulSoup(highlighted, 'html.parser')
            
            # Process each line with proper indentation
            for line_num, line in enumerate(soup.find_all('span', class_='line'), 1):
                # Add line number with proper padding
                num_run = code_para.add_run(f"{line_num:3d} │ ")
                num_run.font.name = style_config["code_font"]
                num_run.font.size = Pt(style_config["code_size"])
                num_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Add code content with syntax highlighting
                for span in line.find_all('span', recursive=False):
                    text = span.get_text()
                    if not text:
                        continue
                        
                    run = code_para.add_run(text)
                    run.font.name = style_config["code_font"]
                    run.font.size = Pt(style_config["code_size"])
                    
                    # Apply token colors
                    color = span.get('style', '').replace('color: ', '')
                    if color:
                        try:
                            if color.startswith('#'):
                                r = int(color[1:3], 16)
                                g = int(color[3:5], 16)
                                b = int(color[5:7], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                        except:
                            pass
                
                code_para.add_run('\n')
            
            # Add border and background
            for run in code_para.runs:
                run._element.rPr.highlight_val = 'lightGray'
            
        except Exception as e:
            logger.error(f"Failed to process code block: {e}")
            # Fallback to simple code block
            p = doc.add_paragraph(code_text, style='Code')

    def _process_paragraph(self, doc: Document, element: Tag, style_config: Dict) -> None:
        """Process a paragraph element and add it to document."""
        try:
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = p.add_run(child.get_text())
                    run.font.name = style_config["code_font"]
                elif child.name == 'a':
                    run = p.add_run(child.get_text())
                    run.font.color.rgb = RGBColor(*style_config["link_color"])
                    run.underline = True
                else:
                    p.add_run(str(child))
        except Exception as e:
            logger.error(f"Failed to process paragraph: {e}")

    def _process_table(self, doc: Document, element: Tag) -> None:
        """Process a table element and add it to document."""
        try:
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
        except Exception as e:
            logger.error(f"Failed to process table: {e}")

    def _process_element(self, doc: Document, element: Tag, style_config: Dict) -> None:
        """Process a single HTML element and add it to document."""
        try:
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            
            elif element.name == 'pre':
                code_block = element.find('code')
                if code_block:
                    # Check if it's a Mermaid diagram
                    if 'language-mermaid' in code_block.get('class', []):
                        code = code_block.get_text().strip()
                        
                        # Add section header
                        doc.add_heading("Mermaid Diagram", level=4)
                        
                        # Add original Mermaid code
                        doc.add_paragraph("Source Code:", style='CodeHeader')
                        self._process_code_block(doc, code_block, style_config)
                        
                        # Generate and add diagram
                        diagram_path, error = self._process_mermaid(code)
                        if diagram_path:
                            # Add diagram with proper styling
                            diagram_para = doc.add_paragraph(style='Diagram')
                            run = diagram_para.add_run()
                            run.add_picture(diagram_path, width=Inches(6.0))
                            
                            # Add caption
                            caption = doc.add_paragraph(style='Caption')
                            caption.add_run("Generated Mermaid Diagram")
                        else:
                            # Add error message
                            error_para = doc.add_paragraph(style='Quote')
                            error_run = error_para.add_run(f"⚠️ Error generating diagram: {error}")
                            error_run.font.color.rgb = RGBColor(255, 0, 0)
                            error_run.bold = True
                    else:
                        self._process_code_block(doc, code_block, style_config)
            
            elif element.name == 'p':
                if element.find('img'):
                    img_src = element.find('img').get('src', '')
                    if img_path := self._handle_image(img_src):
                        self._add_image_to_doc(doc, img_path)
                else:
                    self._process_paragraph(doc, element, style_config)
            
            elif element.name == 'table':
                self._process_table(doc, element)
                
        except Exception as e:
            logger.error(f"Failed to process element {element.name}: {e}")

    def execute(
        self,
        markdown_content: str,
        output_path: str,
        template_path: Optional[str] = None,
        style_config: Optional[str] = None,
    ) -> str:
        """Convert markdown to DOCX format.

        Args:
            markdown_content: Markdown content to convert
            output_path: Output DOCX file path
            template_path: Optional template path
            style_config: Optional style configuration

        Returns:
            Success message with output path

        Raises:
            ValueError: If content is empty or paths invalid
        """
        if not markdown_content.strip():
            raise ValueError("Markdown content cannot be empty")

        try:
            # Setup paths and document
            output_path = self._normalize_path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = self._create_document(template_path)
            style_config = self._parse_style_config(style_config)

            # Convert markdown to HTML
            html = markdown.markdown(
                markdown_content,
                extensions=['fenced_code', 'tables', 'attr_list', 'md_in_html']
            )
            
            # Process elements
            soup = BeautifulSoup(html, 'html.parser')
            for element in soup.find_all():
                self._process_element(doc, element, style_config)

            # Save document
            doc.save(str(output_path))
            return f"Successfully created DOCX: {output_path}"
            
        except Exception as e:
            logger.error(f"Failed to convert markdown: {e}")
            raise ValueError(f"Conversion failed: {str(e)}")


if __name__ == "__main__":
    # Example usage with error handling
    try:
        tool = MarkdownToDocxTool()
        
        # Test markdown with various features
        markdown_content = """
        # Document Title
        
        ## Code Example
        ```python
        def greet(name: str) -> str:
            return f"Hello, {name}!"
        ```
        
        ## System Diagram
        ```mermaid
        graph TD
            A[Start] --> B[Process]
            B --> C[End]
        ```
        
        ## Feature Status
        | Feature | Status |
        |---------|--------|
        | Auth    | Done   |
        | API     | WIP    |
        """
        
        result = tool.execute(
            markdown_content=markdown_content,
            output_path="./example.docx",
            style_config='{"font_name": "Arial"}'
        )
        print(result)
        
    except Exception as e:
        logger.error(f"Example failed: {e}")
